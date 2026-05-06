from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

import tempfile
import unittest

from docx import Document
from scripts.google_docs_builder import create_local_docx


class DocsBuilderTests(unittest.TestCase):
    def _report(self):
        return {
            "title": "Apuntes políticos #7",
            "date_label": "1 de Mayo de 2026",
            "lead": "El oficialismo conserva capacidad de gestión del proceso político.",
            "developments": [
                {
                    "headline": "El Congreso se consolidó como espacio de defensa política",
                    "analysis": "La sesión marcó el principal hecho institucional de la quincena.",
                },
                {
                    "headline": "La calle vuelve a adquirir protagonismo",
                    "analysis": "La protesta sindical gana visibilidad como factor de presión.",
                },
                {
                    "headline": "La opinión pública muestra mayor deterioro",
                    "analysis": "El malestar económico comienza a adquirir carácter acumulativo.",
                },
                {
                    "headline": "El sistema político entra en una fase de movimiento",
                    "analysis": "La oposición muestra reorganización sin alternativa consolidada.",
                },
            ],
            "prospective_keys": [
                "Evolución de la conflictividad social.",
                "Capacidad de control político en Congreso.",
                "Trayectoria de la opinión pública.",
                "Nivel de consolidación opositora.",
            ],
        }

    def test_creates_local_docx_with_apuntes_format(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = create_local_docx(self._report(), report_id="sample", output_dir=Path(tmp))
            self.assertTrue(path.exists())
            doc = Document(str(path))
            full_text = "\n".join(p.text for p in doc.paragraphs)
            table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
            combined = full_text + "\n" + table_text
            self.assertIn("Apuntes políticos #7", combined)
            self.assertIn("NOTA INTERNA", combined)
            self.assertIn("Claves prospectivas", combined)
            self.assertIn("Gracias!", combined)
            self.assertNotIn("BBVA\n", combined)  # logo image replaces text in the brand bar.


if __name__ == "__main__":
    unittest.main()
