from __future__ import annotations

import logging
import os
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from config import REPORTS_OUTPUT_DIR
from google_services import docs_service, drive_service
from utils import parse_bool, split_emails

logger = logging.getLogger(__name__)

ELECTRIC_BLUE = "001391"
SERENE_BLUE = "85C8FF"
SAND = "F7F8F8"
GREY_5 = "000519"
GREY_4 = "46536D"
RED_NOTE = "E60012"
WHITE = "FFFFFF"

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
GOOGLE_DOC_MIME = "application/vnd.google-apps.document"


def _asset_path(*parts: str) -> Path:
    return Path(__file__).resolve().parents[1].joinpath(*parts)


def _rgb(hex_color: str) -> RGBColor:
    hex_color = hex_color.strip().lstrip("#")
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top: int = 0, start: int = 0, bottom: int = 0, end: int = 0) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_no_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def _set_cell_width(cell, width_inches: float) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def _style_run(run, font: str = "Lato", size: float = 10.5, color: str = GREY_5, bold: bool = False, italic: bool = False) -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = _rgb(color)
    run.bold = bold
    run.italic = italic


def _paragraph_spacing(paragraph, before: float = 0, after: float = 0, line: float | None = None) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    if line is not None:
        paragraph.paragraph_format.line_spacing = line


def _set_page_setup(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)


def _content_width(document: Document) -> float:
    section = document.sections[0]
    return (section.page_width - section.left_margin - section.right_margin) / 914400


def _add_brand_bar(document: Document, logo_path: Path) -> None:
    width = _content_width(document)
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_no_borders(table)
    cell = table.cell(0, 0)
    _set_cell_width(cell, width)
    _set_cell_shading(cell, ELECTRIC_BLUE)
    _set_cell_margins(cell, top=170, bottom=145, start=260, end=180)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    row = table.rows[0]
    row.height = Inches(0.58)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    if logo_path.exists():
        run.add_picture(str(logo_path), width=Inches(0.95))
    else:
        run.text = "BBVA"
        _style_run(run, font="Lato", size=16, color=WHITE, bold=True)

    spacer = document.add_paragraph()
    _paragraph_spacing(spacer, after=12)


def _add_meta(document: Document, internal_label: str) -> None:
    width = _content_width(document)
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    _set_no_borders(table)
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    _set_cell_width(left, width * 0.68)
    _set_cell_width(right, width * 0.32)
    for cell in (left, right):
        _set_cell_margins(cell, top=0, bottom=0, start=0, end=0)
    p_left = left.paragraphs[0]
    p_right = right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p_left.add_run("DIRECCIÓN DE RELACIONES INSTITUCIONALES")
    _style_run(run, font="Lato", size=7.8, color=ELECTRIC_BLUE, bold=True)
    run = p_right.add_run(internal_label)
    _style_run(run, font="Lato", size=9.5, color=RED_NOTE, bold=True)
    _paragraph_spacing(p_left, after=4)
    _paragraph_spacing(p_right, after=4)


def _add_title_block(document: Document, report: dict[str, Any]) -> None:
    p = document.add_paragraph()
    _paragraph_spacing(p, before=0, after=1)
    r = p.add_run(str(report.get("title") or "Apuntes políticos"))
    _style_run(r, font="Source Serif 4", size=27, color=ELECTRIC_BLUE, bold=True)

    p = document.add_paragraph()
    _paragraph_spacing(p, before=0, after=18)
    r = p.add_run(str(report.get("date_label") or ""))
    _style_run(r, font="Lato", size=10.5, color=GREY_5)

    p = document.add_paragraph()
    _paragraph_spacing(p, before=5, after=18, line=0.94)
    r = p.add_run(str(report.get("lead") or ""))
    _style_run(r, font="Source Serif 4", size=12.8, color=ELECTRIC_BLUE, bold=True)


def _add_developments_block(document: Document, report: dict[str, Any]) -> None:
    width = _content_width(document)
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_no_borders(table)
    cell = table.cell(0, 0)
    _set_cell_width(cell, width)
    _set_cell_shading(cell, SAND)
    _set_cell_margins(cell, top=220, bottom=170, start=300, end=260)

    # Remove initial empty paragraph after using it for the first item.
    first = True
    for item in report.get("developments", []):
        paragraph = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        _paragraph_spacing(paragraph, before=0, after=10, line=1.04)
        paragraph.paragraph_format.left_indent = Inches(0.24)
        paragraph.paragraph_format.first_line_indent = Inches(-0.20)

        bullet = paragraph.add_run("─ ")
        _style_run(bullet, font="Lato", size=10.3, color=GREY_4)

        headline = str(item.get("headline") or "").strip()
        if headline and not headline.endswith((".", "?", "!")):
            headline += "."
        analysis = str(item.get("analysis") or "").strip()
        h = paragraph.add_run(headline + (" " if analysis else ""))
        _style_run(h, font="Lato", size=10.3, color=ELECTRIC_BLUE, bold=True)
        a = paragraph.add_run(analysis)
        _style_run(a, font="Lato", size=10.3, color=GREY_5)

    keys_heading = cell.add_paragraph()
    _paragraph_spacing(keys_heading, before=6, after=5, line=1.0)
    keys_heading.paragraph_format.left_indent = Inches(0.24)
    keys_heading.paragraph_format.first_line_indent = Inches(-0.20)
    b = keys_heading.add_run("─ ")
    _style_run(b, font="Lato", size=10.4, color=GREY_4)
    r = keys_heading.add_run("Claves prospectivas")
    _style_run(r, font="Lato", size=10.4, color=ELECTRIC_BLUE, bold=True)

    for key in report.get("prospective_keys", []):
        p = cell.add_paragraph()
        _paragraph_spacing(p, before=0, after=1, line=0.96)
        p.paragraph_format.left_indent = Inches(0.58)
        p.paragraph_format.first_line_indent = Inches(-0.16)
        r = p.add_run("○ ")
        _style_run(r, font="Lato", size=9.5, color=ELECTRIC_BLUE)
        r = p.add_run(str(key).strip())
        _style_run(r, font="Lato", size=9.5, color=ELECTRIC_BLUE, bold=True)


def _add_footer(document: Document) -> None:
    spacer = document.add_paragraph()
    _paragraph_spacing(spacer, after=14)
    width = _content_width(document)
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_no_borders(table)
    cell = table.cell(0, 0)
    _set_cell_width(cell, width)
    _set_cell_shading(cell, SERENE_BLUE)
    _set_cell_margins(cell, top=100, bottom=95, start=0, end=0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _paragraph_spacing(p, before=0, after=0)
    r = p.add_run("Gracias!")
    _style_run(r, font="Lato", size=11, color=WHITE, bold=True)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _paragraph_spacing(p, before=2, after=0)
    r = p.add_run("DIRECCIÓN DE RELACIONES INSTITUCIONALES")
    _style_run(r, font="Lato", size=7, color=ELECTRIC_BLUE, bold=True)


def create_local_docx(report: dict[str, Any], report_id: str, output_dir: Path | None = None) -> Path:
    output_dir = output_dir or (REPORTS_OUTPUT_DIR / report_id)
    output_dir.mkdir(parents=True, exist_ok=True)
    docx_path = output_dir / f"{report_id}.docx"
    document = Document()
    _set_page_setup(document)
    logo_path = _asset_path("assets", "brand", "logo_bbva_white.png")
    _add_brand_bar(document, logo_path)
    _add_meta(document, os.getenv("REPORT_INTERNAL_LABEL", "NOTA INTERNA"))
    _add_title_block(document, report)
    _add_developments_block(document, report)
    _add_footer(document)
    document.save(docx_path)
    return docx_path


def _share_file(drive, file_id: str) -> list[dict[str, Any]]:
    actions: list[dict[str, Any]] = []
    share_with = split_emails(os.getenv("GOOGLE_DOCS_SHARE_WITH", ""))
    role = os.getenv("GOOGLE_DOCS_SHARE_ROLE", "writer").strip() or "writer"
    for email in share_with:
        drive.permissions().create(
            fileId=file_id,
            body={"type": "user", "role": role, "emailAddress": email},
            sendNotificationEmail=False,
            supportsAllDrives=True,
        ).execute()
        actions.append({"email": email, "role": role, "status": "shared"})

    if parse_bool(os.getenv("GOOGLE_DOCS_ANYONE_WITH_LINK"), False):
        drive.permissions().create(
            fileId=file_id,
            body={"type": "anyone", "role": "reader"},
            sendNotificationEmail=False,
            supportsAllDrives=True,
        ).execute()
        actions.append({"type": "anyone", "role": "reader", "status": "shared"})
    return actions


def _try_transfer_ownership(drive, file_id: str) -> list[dict[str, Any]]:
    results: list[dict[str, Any]] = []
    transfer_to = split_emails(os.getenv("GOOGLE_DOCS_TRANSFER_OWNERSHIP_TO", ""))
    for email in transfer_to:
        try:
            drive.permissions().create(
                fileId=file_id,
                body={"type": "user", "role": "owner", "emailAddress": email},
                transferOwnership=True,
                sendNotificationEmail=True,
                supportsAllDrives=True,
            ).execute()
            results.append({"email": email, "role": "owner", "status": "transfer_requested"})
        except Exception as exc:  # noqa: BLE001 - ownership transfer should not break report delivery.
            logger.warning("event=ownership_transfer_failed email=%s reason=%s", email, exc)
            results.append({"email": email, "role": "owner", "status": "failed", "reason": str(exc)[:300]})
    return results


def upload_docx_as_google_doc(docx_path: Path, title: str) -> dict[str, Any]:
    from googleapiclient.http import MediaFileUpload

    drive = drive_service()
    metadata: dict[str, Any] = {"name": title, "mimeType": GOOGLE_DOC_MIME}
    folder_id = os.getenv("GOOGLE_DOCS_FOLDER_ID", "").strip()
    if folder_id:
        metadata["parents"] = [folder_id]
    media = MediaFileUpload(str(docx_path), mimetype=DOCX_MIME, resumable=False)
    created = (
        drive.files()
        .create(
            body=metadata,
            media_body=media,
            fields="id,name,webViewLink,mimeType,parents",
            supportsAllDrives=True,
        )
        .execute()
    )
    file_id = created["id"]
    sharing = _share_file(drive, file_id)
    ownership = _try_transfer_ownership(drive, file_id)

    # Validate the converted file through the Docs API. The document was authored locally
    # as DOCX, uploaded with Drive API and converted to a native Google Doc.
    try:
        docs_service().documents().get(documentId=file_id, fields="documentId,title").execute()
    except Exception as exc:  # noqa: BLE001
        logger.warning("event=docs_validation_failed file_id=%s reason=%s", file_id, exc)

    meta = drive.files().get(fileId=file_id, fields="id,name,webViewLink,mimeType", supportsAllDrives=True).execute()
    return {
        "document_id": file_id,
        "document_url": meta.get("webViewLink") or f"https://docs.google.com/document/d/{file_id}/edit",
        "name": meta.get("name", title),
        "mime_type": meta.get("mimeType", GOOGLE_DOC_MIME),
        "local_docx_path": str(docx_path),
        "sharing": sharing,
        "ownership_transfer": ownership,
    }


def create_google_doc(report: dict[str, Any], report_id: str) -> dict[str, Any]:
    title = f"{report['title']} - {report.get('date_label', '')}".strip(" -")
    out_dir = REPORTS_OUTPUT_DIR / report_id
    docx_path = create_local_docx(report, report_id=report_id, output_dir=out_dir)
    return upload_docx_as_google_doc(docx_path, title=title)
